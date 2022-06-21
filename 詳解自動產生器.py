#!/usr/bin/env python
# coding: utf-8

# <ul>
# 目前有的特殊前綴格式：
#     <li>*ph* = 片語</li>
#     <li>*tip* = 註解</li>
# </ul>
# <p>有些解釋後面接著"；"，那其以後的解釋(如果存在)便會直接接上</p>
# <p>如果沒有"；"，那其以後的解釋(如果存在)便會換到下一行</p>

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
import time

wordsFilePath = 'C:\\Users\\User\\OneDrive\\Documents\\文林心相關\\詳解\\詳解們\\107-2全模(第2回)\\Voc.txt'
savedFileName = None
dicName = 'C:\\Users\\User\\OneDrive\\Documents\\文林心相關\\詳解\\詳解需求\\字典.txt'
partsOfSpeech = ['(n.)', '(a.)', '(adj.)', '(adv.)', '(v.)', '(vi.)', '(vt.)', '(prep.)', '(conj.)', '(pron.)', '(int.)', '(aux.)', '(pl.), ']
urlForSearch = 'https://tw.dictionary.search.yahoo.com/search;_ylt=AwrtXWrsmTNfJ2AAQj17rolQ?p='

def cutTabMinimizeSpace(line):
    
    length = len(line)
    i = 0
    while True:
        if i >= length:
            break
        if line[i] == '\t':
            line = line[0:i] + line[i + 1:]
            length -= 1
        elif line[i] == ' ':
            if line[i + 1] != ' ':
                i += 1
            else:
                line = line[0:i] + line[i + 1:]
                length -= 1
        else:
            i += 1
    if line[0] == ' ': 
        return line[1:]
    else:
        return line
    
def cutNumber(line):
    
    while line[0].isnumeric():
        line = line[1:]
    return line[1:]

def cutHeadAndTailSpace(line):
    
    string = line
    while string[-1] == ' ':
        string = string[:len(string) - 1]
    while string[0] == ' ':
        string = string[1:]
    while string[0] == '\t':
        string = string[1:]
    while string[-1] == '\t':
        string = string[:len(string) - 1]
    while string[0] == '\u3000':
        string = string[1:]
    while string[-1] == '\u3000':
        string = string[:len(string) - 1]
    return string

def getDic(dicName):

    f = open(dicName, 'r', encoding = 'utf-16')

    lines = []
    dic = []
    group = []

    for line in f:
        tempStr = cutTabMinimizeSpace(line)[:-1]
        if len(tempStr) is not 0:
            lines.append(cutTabMinimizeSpace(line)[:-1])

    group.append(cutNumber(lines[0]))
    for line in lines[1:]:
        if line[0] >= '0' and line[0] <= '9':
            dic.append(group)
            group = []
            group.append(cutHeadAndTailSpace(cutNumber(line)))
        else:
            group.append(cutHeadAndTailSpace(line))
    dic.append(group)

    for groupNum in range(0, len(dic)):
        i = 0
        length = len(dic[groupNum])
        while True:
            if i + 1 >= length:
                break
            if dic[groupNum][i][-1] == '；':
                canUse = True
                for string in partsOfSpeech:
                    if dic[groupNum][i + 1].find(string) is not -1:
                        canUse = False
                        break
                if canUse:
                    dic[groupNum][i] += dic[groupNum][i + 1]
                    dic[groupNum] = dic[groupNum][0:i + 1] + dic[groupNum][i + 2:]
                    length -= 1
                else:
                    i += 1
            else:
                canUse = True
                for string in partsOfSpeech:
                    if dic[groupNum][i + 1].find(string) is not -1:
                        canUse = False
                        break
                if dic[groupNum][i + 1][0] != '=' and dic[groupNum][i + 1][0] != '↔' and dic[groupNum][i + 1][0] != '(':
                    canUse = False
                if canUse:
                    dic[groupNum][i] += dic[groupNum][i + 1]
                    dic[groupNum] = dic[groupNum][0:i + 1] + dic[groupNum][i + 2:]
                    length -= 1
                else:
                    i += 1

    for group in dic:
        print(group)
    print('--------------------------------------------------------------------------------')
    print('--------------------------------------------------------------------------------')
    print('--------------------------------------------------------------------------------')

    return dic

def searchForWordInGroup(word):
    for group in dic:
        for line in group:
            location = line.find(word)
            if location == 0 and line[len(word)] == ' ':
                return group
    return None

def getVoc(line):
    return line.split(' ')[0]
    
def getPartsOfSpeech(line):
    string = line
    returnStr = ''
    while True:
        found = -1
        for part in partsOfSpeech:
            found = string.find(part)
            if found is not -1:
                break
        if found is not -1:
            if found is not 0:
                string = string[found:]
            returnStr += string.split(')')[0] + ')'
            string = string[(len(string.split(')')[0]) + 1):]
        else:
            break
            
    return returnStr
    
def getExplaination(line):
    string = str(line)
    while True:
        found = -1
        for part in partsOfSpeech:
            found = string.find(part)
            if found is not -1:
                break
        if found is not -1:
            if found is not 0:
                string = string[found:]
            string = string[(len(string.split(')')[0]) + 1):]
        else:
            break
    
    return string
    
def is_all_chinese(strs):
    for _char in strs:
        if not '\u4e00' <= _char <= '\u9fa5':
            return False
    return True

def addParen(line):
    line = '(' + line
    i = 0
    length = len(line)
    parenNum = 1
    while True:
        if (i >= length):
            break
        if line[i] == '.' and parenNum >= 1:
            line = line[0:i + 1] + ')' + line[i + 1:]
            i += 2
            length += 1
            parenNum -= 1
            if line[i].isalpha() and not is_all_chinese(line[i]):
                print(line[i])
                line = line[0:i] + '(' + line[i:]
                parenNum += 1
                i += 1
                length += 1
        else:
            i += 1
    return line

def searchFromYahoo(voc):
       
    resp = requests.get(urlForSearch + voc)

    resp.encoding = 'utf-8'    #轉換編碼至UTF-8
    
    soup = BeautifulSoup(resp.text, "lxml")
    if soup.find(text = '很抱歉，字典找不到您要的資料喔！') != None:
        return [voc, 'NOT', 'FOUND']
    
    group = []
    
    lines = soup.find(class_="dd cardDesign dictionaryWordCard sys_dict_word_card").contents[1].find(class_="compList mb-25 p-rel").contents[0].contents
    
    for i in range(0, len(lines)):
        line = lines[i]
        if i == 0:
            group.append(voc + ' ' + addParen(line.contents[0].text + line.contents[2].text))
        else:
            group.append(addParen(line.contents[0].text + line.contents[2].text))
    return group

def makeVocWord(dic, filePath, wordName = None):
    
    if wordName == None:
        wordName = time.strftime("%Y%m%d_%H%M%S", time.localtime()) + '.docx'
        
    if len(wordName.split('.')) == 1:
        wordName += '.docx'
    elif wordName.split('.')[-1] != 'docx' and wordName.split('.')[-1] != 'doc':
        wordName = wordName[:(len(wordName) - len(wordName.split('.')[-1]) + 1)] + '.docx'

    f = open(filePath, 'r')

    vocList = f.readlines()

    document = Document()

    table = document.add_table(rows = 1, cols = 4)
    table.style = 'TableGrid'

    currentRow = table.rows[0]

    currentRow.cells[0].text = '題號'
    currentRow.cells[1].text = '單字'
    currentRow.cells[2].text = '詞性'
    currentRow.cells[3].text = '解釋'

    table.add_row()
    currentRow = table.rows[-1]
    currentRow.cells[0].text = '1.'

    i = 1

    sectionStartRow = 0
    vocStartRow = 0

    totalNum = 0
    for voc in vocList:
        if voc != '\n':
            totalNum += 1

    currentNum = 0
    for voc in vocList:

        if voc != '\n':
            currentNum += 1
        else:
            i += 1
            currentRow.cells[0].text = str(i) + '.'
            continue

        print('Processing ' + str(currentNum) + '/' + str(totalNum) + ' ' + voc)

        if voc[-1] == '\n':
            group = searchForWordInGroup(voc[:-1])
        else:
            group = searchForWordInGroup(voc)

        if group == None:
            group = searchFromYahoo(voc)

        if 'NOT' in group and 'FOUND' in group:
            currentRow.cells[1].text = group[0]
            currentRow.cells[2].text = group[1]
            currentRow.cells[3].text = group[2]
        else: 
            for string in group:

                if string[0] == '(':
                    currentRow.cells[2].text = getPartsOfSpeech(string)
                    currentRow.cells[3].text = getExplaination(string)
                elif string[0] == '*':
                    currentRow.cells[1].text = string
                elif string[0].isalpha() and not is_all_chinese(string[0]):
                    currentRow.cells[1].text = getVoc(string)
                    currentRow.cells[2].text = getPartsOfSpeech(string)
                    currentRow.cells[3].text = getExplaination(string)
                else:
                    currentRow.cells[3].text = string

                table.add_row()
                currentRow = table.rows[-1]      

    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(14)
                    font.name = 'Times New Roman'
    f.close()
    document.save(wordName)
    print('Successfully save ' + wordName)

dic = getDic()
makeVocWord(dic, wordsFilePath, savedFileName)
